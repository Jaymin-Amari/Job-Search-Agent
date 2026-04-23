#!/usr/bin/env python3
"""
Job Search Agent — main orchestration loop.

Run:
  python agent.py               # live run (reads/writes Drive)
  DRY_RUN=true python agent.py  # prints briefing to stdout, no Drive writes
"""

import io
import json
import os
import re
import sys
from datetime import datetime, timezone

import anthropic
from docx import Document

import config
import drive_handler
import prompts
from sources import gmail_linkedin, indeed, scraper


def main() -> None:
    dry_run = config.DRY_RUN
    if dry_run:
        print("[DRY RUN] Drive writes disabled — briefing will print to stdout.")

    client = anthropic.Anthropic()
    today = datetime.now(timezone.utc).strftime("%B %d, %Y").upper()
    date_prefix = datetime.now(timezone.utc).strftime("%Y-%m-%d")

    # ── 1. Load persistent state from Drive ───────────────────────────────────
    print("Reading master resume from Drive...")
    resume_text = drive_handler.download_docx_text(config.MASTER_RESUME_ID)

    print("Reading seen_jobs.log from Drive...")
    seen_raw = drive_handler.read_text_file_in_folder(
        config.DRIVE_OUTPUT_FOLDER_ID, config.SEEN_JOBS_FILE_NAME
    )
    seen_urls: set[str] = {line.strip() for line in seen_raw.splitlines() if line.strip()}

    # ── 2. Fetch all jobs ──────────────────────────────────────────────────────
    print("Fetching LinkedIn jobs from Drive staging file...")
    linkedin_jobs = gmail_linkedin.get_linkedin_jobs(dry_run=dry_run)

    print("Searching Indeed via MCP...")
    indeed_jobs = indeed.search_indeed(client)

    print("Scraping watchlist career pages...")
    watchlist_jobs, scraper_notes = scraper.get_watchlist_jobs()

    all_jobs = linkedin_jobs + indeed_jobs + watchlist_jobs
    print(f"Total jobs fetched: {len(all_jobs)}")

    # ── 3. Deduplicate against seen_jobs.log ──────────────────────────────────
    new_jobs: list[dict] = []
    for job in all_jobs:
        url = job.get("url", "").strip()
        if url and url not in seen_urls:
            new_jobs.append(job)
            seen_urls.add(url)  # prevent dupes within this run too

    print(f"New (unseen): {len(new_jobs)}")

    # ── 4. Pre-filter: excluded titles ────────────────────────────────────────
    def _is_excluded(title: str) -> bool:
        t = title.lower()
        return any(ex.lower() in t for ex in config.EXCLUDED_TITLES)

    jobs_to_score = [j for j in new_jobs if not _is_excluded(j.get("title", ""))]
    pre_filtered = len(new_jobs) - len(jobs_to_score)
    if pre_filtered:
        print(f"Pre-filtered (excluded titles): {pre_filtered}")

    # ── 5. Score and process each job ─────────────────────────────────────────
    strong_fits: list[dict] = []
    worth_a_look: list[dict] = []
    skipped: list[dict] = []
    gated_out: list[dict] = []

    for job in jobs_to_score:
        title = job.get("title", "Unknown Role")
        company = job.get("company", "Unknown Company")
        print(f"  Scoring: {title} — {company}")

        # Fetch description if scraper deferred it
        description = job.get("description", "")
        if not description and job.get("url"):
            description = _fetch_description(job["url"])
            job["description"] = description

        job_text = (
            f"Title: {title}\n"
            f"Company: {company}\n"
            f"Location: {job.get('location', '')}\n\n"
            f"{description}"
        )

        # Flag queer-focused roles for manual handling
        if _is_queer_focused(job_text):
            print(f"    ⚑ Queer-focused role — flagged for manual handling, no auto-apply")
            strong_fits.append({**job, "_queer_flag": True})
            continue

        # Score via Claude API
        score_result = _score_job(client, resume_text, job_text, job.get("url", ""))
        if score_result is None:
            skipped.append(job)
            continue

        score = score_result.get("score", 0)
        gate_passed = score_result.get("gate_passed", False)

        if not gate_passed:
            gated_out.append({**job, **score_result})
            continue

        network_flag = score_result.get("network_flag") or _check_network_flag(company)

        if score >= 7:
            print(f"    Score {score} — gap analysis + writing...")
            gap = _run_gap_analysis(client, resume_text, job_text)
            resume_bytes = _tailor_resume(client, resume_text, job_text, gap)
            cover_bytes = _write_cover_letter(client, resume_text, job_text, gap)

            safe_co = _safe_name(company)
            safe_title = _safe_name(title)
            resume_fname = f"{date_prefix} — {safe_co} — {safe_title} — Resume.docx"
            cover_fname = f"{date_prefix} — {safe_co} — {safe_title} — Cover Letter.docx"

            if not dry_run:
                drive_handler.upload_docx(config.DRIVE_OUTPUT_FOLDER_ID, resume_fname, resume_bytes)
                drive_handler.upload_docx(config.DRIVE_OUTPUT_FOLDER_ID, cover_fname, cover_bytes)
                print(f"    Uploaded: {resume_fname}")
                print(f"    Uploaded: {cover_fname}")
            else:
                print(f"    [DRY RUN] Would upload: {resume_fname}")
                print(f"    [DRY RUN] Would upload: {cover_fname}")

            strong_fits.append({**job, **score_result, "network_flag": network_flag})

        elif score >= 5:
            worth_a_look.append({**job, **score_result, "network_flag": network_flag})
        else:
            skipped.append({**job, **score_result})

    # ── 6. Build briefing ──────────────────────────────────────────────────────
    scored_count = len(jobs_to_score) - len(gated_out)
    strong_real = [j for j in strong_fits if not j.get("_queer_flag")]

    stats = {
        "sources_scanned": (
            f"{len(config.TARGET_ROLES)} board searches + "
            f"{len(config.WATCHLIST_CONFIRMED) + len(config.WATCHLIST_AUTO_FIND)} watchlist pages"
        ),
        "total_reviewed": len(jobs_to_score),
        "gated_out": len(gated_out),
        "scored": scored_count,
        "strong_fits": len(strong_real),
        "worth_a_look": len(worth_a_look),
        "skipped": len(skipped),
    }

    briefing = _format_briefing(today, stats, strong_fits, worth_a_look, skipped, scraper_notes)

    # ── 7. Write outputs ───────────────────────────────────────────────────────
    if dry_run:
        print("\n" + "=" * 60)
        print(briefing)
        print("=" * 60)
    else:
        print("Updating Daily Briefing Google Doc...")
        doc_id = drive_handler.get_or_create_briefing_doc(
            config.DRIVE_OUTPUT_FOLDER_ID, config.DAILY_BRIEFING_DOC_NAME
        )
        drive_handler.prepend_to_doc(doc_id, briefing)

        print("Writing seen_jobs.log back to Drive...")
        drive_handler.write_text_file_in_folder(
            config.DRIVE_OUTPUT_FOLDER_ID,
            config.SEEN_JOBS_FILE_NAME,
            "\n".join(seen_urls),
        )

    print(
        f"\nDone. {len(strong_real)} strong fit(s), "
        f"{len(worth_a_look)} worth a look, "
        f"{len(gated_out)} gated out."
    )


# ── Claude API calls ───────────────────────────────────────────────────────────

def _score_job(
    client: anthropic.Anthropic,
    resume_text: str,
    job_text: str,
    job_url: str,
) -> dict | None:
    full_prompt = prompts.SCORING_PROMPT.format(resume_text=resume_text, job_text=job_text)

    # Split at "Job posting:" to cache the resume + instructions portion across calls
    split_at = "Job posting:"
    if split_at in full_prompt:
        idx = full_prompt.index(split_at)
        cacheable_part = full_prompt[:idx]
        variable_part = full_prompt[idx:]
        content: list | str = [
            {
                "type": "text",
                "text": cacheable_part,
                "cache_control": {"type": "ephemeral"},
            },
            {"type": "text", "text": variable_part},
        ]
    else:
        content = full_prompt

    try:
        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=1024,
            messages=[{"role": "user", "content": content}],
            extra_headers={"anthropic-beta": "prompt-caching-2024-07-31"},
        )
        text = response.content[0].text.strip()
        start = text.find("{")
        end = text.rfind("}") + 1
        if start >= 0 and end > start:
            result = json.loads(text[start:end])
            # Inject the actual URL into the teal_entry if missing
            if "teal_entry" in result and not result["teal_entry"].get("url"):
                result["teal_entry"]["url"] = job_url
            return result
    except Exception as e:
        print(f"    Scoring error: {e}")
    return None


def _run_gap_analysis(
    client: anthropic.Anthropic, resume_text: str, job_text: str
) -> str:
    prompt = prompts.GAP_ANALYSIS_PROMPT.format(resume_text=resume_text, job_text=job_text)
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )
    return response.content[0].text


def _tailor_resume(
    client: anthropic.Anthropic,
    resume_text: str,
    job_text: str,
    gap: str,
) -> bytes:
    prompt = prompts.RESUME_TAILORING_PROMPT.format(
        resume_text=resume_text,
        job_text=job_text,
        gap_analysis_output=gap,
    )
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )
    return _text_to_docx(response.content[0].text)


def _write_cover_letter(
    client: anthropic.Anthropic,
    resume_text: str,
    job_text: str,
    gap: str,
) -> bytes:
    prompt = prompts.COVER_LETTER_PROMPT.format(
        resume_text=resume_text,
        job_text=job_text,
        gap_analysis_output=gap,
    )
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1024,
        messages=[{"role": "user", "content": prompt}],
    )
    return _text_to_docx(response.content[0].text)


# ── Document generation ────────────────────────────────────────────────────────

def _text_to_docx(text: str) -> bytes:
    """Convert plain text (with basic Markdown hints) to a Word doc."""
    doc = Document()
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith(("- ", "• ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped:
            doc.add_paragraph(stripped)
        else:
            doc.add_paragraph("")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Briefing formatter ─────────────────────────────────────────────────────────

def _format_briefing(
    today: str,
    stats: dict,
    strong_fits: list[dict],
    worth_a_look: list[dict],
    skipped: list[dict],
    scraper_notes: list[str],
) -> str:
    lines = [
        "═══════════════════════════════════════",
        f"📅 {today} — DAILY JOB BRIEFING",
        f"Scanned: {stats['sources_scanned']}",
        (
            f"New postings reviewed: {stats['total_reviewed']} | "
            f"Gated out: {stats['gated_out']} | "
            f"Scored: {stats['scored']}"
        ),
        (
            f"Briefing includes: {stats['strong_fits']} strong fits | "
            f"{stats['worth_a_look']} worth a look | "
            f"{stats['skipped']} skipped"
        ),
        "═══════════════════════════════════════",
        "",
        "── 🟢 STRONG FITS (7+) ──────────────────",
        "",
    ]

    for job in strong_fits:
        title = job.get("title", "Unknown Role")
        company = job.get("company", "Unknown Company")
        url = job.get("url", "")
        teal = job.get("teal_entry") or {}

        if job.get("_queer_flag"):
            lines += [
                f"### {title} — {company}  ⚑ QUEER-FOCUSED ROLE",
                "Manual handling required — queer context may apply. Do not auto-apply.",
                f"🔗 {url}",
                f"📌 {title} — {company}",
                "",
                "---",
                "",
            ]
            continue

        score = job.get("score", "N/A")
        fit_reasons = job.get("fit_reasons", [])
        concerns = job.get("concerns", [])

        lines += [
            f"### {title} — {company} (Score: {score}/10)",
            f"Summary: {job.get('fit_summary', '')}",
            f"Fit: {'; '.join(fit_reasons)}",
            f"Watch: {'; '.join(concerns)}",
            f"Interview likelihood: {job.get('interview_likelihood', '')}",
            f"Offer likelihood: {job.get('offer_likelihood', '')}",
        ]

        if job.get("network_flag"):
            lines.append(f"⚑ Network: {job['network_flag']}")

        display_title = teal.get("title") or title
        display_company = teal.get("company") or company
        lines += [
            f"🔗 {url}",
            f"📌 {display_title} — {display_company}",
            "📁 Resume + Cover Letter: saved to Drive output folder",
            "",
            "---",
            "",
        ]

    lines += [
        "── 🟡 WORTH A LOOK (5–6) — No draft generated ──",
        "",
    ]

    for job in worth_a_look:
        score = job.get("score", "N/A")
        title = job.get("title", "Unknown Role")
        company = job.get("company", "Unknown Company")
        concerns = job.get("concerns", [])
        note = concerns[0] if concerns else ""
        salary_flag = " (salary unconfirmed)" if any(
            w in note.lower() for w in ["salary", "compensation", "pay"]
        ) else ""
        network = f" | ⚑ {job['network_flag']}" if job.get("network_flag") else ""
        lines.append(f"- {title}, {company} ({score}) — {note}{salary_flag}{network}")

    lines += [
        "",
        f"── ⚫ SKIPPED (<5): {len(skipped)} postings ─────────────",
    ]

    if scraper_notes:
        lines += [
            "",
            "── ⚠ SCRAPER NOTES ──────────────────────",
        ]
        for note in scraper_notes:
            lines.append(f"  • {note}")

    lines.append("")
    return "\n".join(lines)


# ── Utility helpers ────────────────────────────────────────────────────────────

def _fetch_description(url: str) -> str:
    try:
        import requests
        from bs4 import BeautifulSoup
        headers = {"User-Agent": "Mozilla/5.0 (compatible; JobAgent/1.0)"}
        resp = requests.get(url, headers=headers, timeout=15)
        if resp.status_code != 200:
            return ""
        soup = BeautifulSoup(resp.text, "lxml")
        for tag in soup(["script", "style", "nav", "header", "footer"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)[:5000]
    except Exception:
        return ""


def _is_queer_focused(text: str) -> bool:
    t = text.lower()
    return any(kw in t for kw in config.QUEER_ROLE_KEYWORDS)


def _check_network_flag(company: str) -> str | None:
    company_lower = company.lower()
    for key, note in config.NETWORK_FLAGS.items():
        if key in company_lower:
            return note
    return None


def _safe_name(s: str) -> str:
    """Strip characters illegal in filenames and truncate."""
    return re.sub(r'[\\/*?:"<>|]', "", s).strip()[:40]


if __name__ == "__main__":
    main()
